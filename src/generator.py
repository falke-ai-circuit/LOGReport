from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Preformatted, PageBreak
from reportlab.platypus.tableofcontents import TableOfContents
from docx import Document
from docx.shared import Pt
from typing import List, Dict
from pathlib import Path
import os
import re
import textwrap

# Handle both runtime (relative) and test (absolute) import contexts
try:
    from utils.file_utils import filter_lines  # Runtime from src/
except ImportError:
    from src.utils.file_utils import filter_lines  # Test context from project root

class ReportGenerator:
    # File type processing order for node grouping
    TYPE_ORDER = {'.fbc': 0, '.rpc': 1, '.log': 2, '.lis': 3}
    
    def __init__(self):
        self.styles = {
            'pdf': {
                'title': ParagraphStyle(
                    'Title',
                    fontName='Helvetica-Bold',
                    fontSize=14,
                    spaceAfter=14,
                    textColor='#5D3E8E'
                ),
                'subtitle': ParagraphStyle(
                    'Subtitle',
                    fontName='Helvetica-Bold',
                    fontSize=12,
                    spaceAfter=6,
                    textColor='#7A5299'
                ),
                'body': ParagraphStyle(
                    'Body',
                    fontName='Courier',
                    fontSize=10,
                    leading=12
                ),
                'node_chapter': ParagraphStyle(
                    'NodeChapter',
                    fontName='Helvetica-Bold',
                    fontSize=16,
                    spaceAfter=10,
                    textColor='#4A148C',
                    keepWithNext=True
                ),
                'file_subheading': ParagraphStyle(
                    'FileSubheading',
                    fontName='Helvetica-Bold',
                    fontSize=11,
                    spaceAfter=4,
                    textColor='#6A1B9A'
                )
            }
        }

    def extract_node_from_filename(self, filename: str) -> str:
        """Extract node name from filename (e.g., AP01m, AL02)"""
        # Match patterns like AP01m, AP02r, AL01, etc.
        match = re.match(r'(AP\d{2}[mr]?|AL\d{2})', filename, re.IGNORECASE)
        if match:
            return match.group(1).upper()
        return "Unknown"
    
    def group_logs_by_node(self, logs: List[Dict]) -> Dict[str, Dict[str, List[Dict]]]:
        """Group logs by node name and file type"""
        grouped = {}
        for log in logs:
            node_name = self.extract_node_from_filename(log['filename'])
            if node_name not in grouped:
                grouped[node_name] = {'.fbc': [], '.rpc': [], '.log': [], '.lis': []}
            
            # Determine file type from extension
            filename_lower = log['filename'].lower()
            file_type = None
            for ext in self.TYPE_ORDER.keys():
                if ext in filename_lower:
                    file_type = ext
                    break
            
            if file_type:
                grouped[node_name][file_type].append(log)
            else:
                # Fallback to .log for unknown types
                grouped[node_name]['.log'].append(log)
        
        return grouped
    
    def wrap_long_lines(self, content: List[str], max_width: int = None) -> List[str]:
        """Wrap long lines to prevent text cutoff on page
        
        Args:
            content: List of text lines to wrap
            max_width: Maximum characters per line. If None, calculates based on A4 page
                      with 20mm margins and Courier 10pt font (80 chars fits perfectly)
        """
        if max_width is None:
            # A4 width: 210mm, left+right margins: 40mm, usable: 170mm
            # Courier 10pt: 80 chars = 480 points = 169.33mm (verified with stringWidth)
            max_width = 80
            
        wrapped_lines = []
        for line in content:
            if len(line) > max_width:
                # Use textwrap to split at word boundaries, fallback to character split
                wrapped = textwrap.wrap(line, width=max_width, 
                                       break_long_words=True,  # Allow breaking long words
                                       break_on_hyphens=False)
                wrapped_lines.extend(wrapped)
            else:
                wrapped_lines.append(line)
        return wrapped_lines

    def generate_pdf(self, logs: List[Dict], output_path: str, 
                    lines_mode: str = 'all', 
                    line_limit: int = 0,
                    range_start: int = 0, 
                    range_end: int = 0):
        """Generate PDF with node-based grouping, bookmarks, and line wrapping"""
        try:
            output_path = str(Path(output_path).absolute())
            if not output_path.lower().endswith('.pdf'):
                output_path += '.pdf'
                
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                leftMargin=20*mm,
                rightMargin=20*mm,
                topMargin=20*mm,
                bottomMargin=20*mm
            )
            
            story = []
            
            # Group logs by node
            grouped_logs = self.group_logs_by_node(logs)
            
            # Add Table of Contents at the beginning
            story.append(Paragraph('Log Report - Table of Contents', self.styles['pdf']['title']))
            story.append(Spacer(1, 10*mm))
            
            for node_name in sorted(grouped_logs.keys()):
                # Create clickable link to node section
                toc_entry = f'<a href="#{node_name}" color="blue"><u>Node: {node_name}</u></a>'
                story.append(Paragraph(toc_entry, self.styles['pdf']['subtitle']))
                story.append(Spacer(1, 2*mm))
            
            story.append(PageBreak())
            
            # Sort nodes alphabetically
            for node_name in sorted(grouped_logs.keys()):
                # Add node chapter with anchor bookmark (using <a> tag for PDF navigation)
                chapter_text = f'<a name="{node_name}"/>Node: {node_name}'
                story.append(Paragraph(
                    chapter_text,
                    self.styles['pdf']['node_chapter']
                ))
                story.append(Spacer(1, 6*mm))
                
                # Process file types in order (.fbc → .rpc → .log → .lis)
                for file_type in sorted(grouped_logs[node_name].keys(), 
                                       key=lambda x: self.TYPE_ORDER.get(x, 999)):
                    file_logs = grouped_logs[node_name][file_type]
                    
                    if not file_logs:  # Skip empty file types
                        continue
                    
                    for log in file_logs:
                        # Create a copy of the log to modify
                        processed_log = dict(log)
                        
                        # Apply line filtering based on parameters
                        processed_log['content'] = filter_lines(
                            processed_log['content'],
                            mode=lines_mode,
                            limit=line_limit,
                            start=range_start,
                            end=range_end
                        )
                        
                        # Wrap long lines to prevent cutoff
                        processed_log['content'] = self.wrap_long_lines(processed_log['content'])
                        
                        # Add file subheading
                        story.append(Paragraph(
                            f"{file_type.upper()[1:]} File: {processed_log['filename']}",
                            self.styles['pdf']['file_subheading']
                        ))
                        
                        # Add filtered and wrapped content
                        content_text = "\n".join(processed_log['content'])
                        story.append(Preformatted(content_text, self.styles['pdf']['body']))
                        
                        story.append(Spacer(1, 6*mm))
                
                # Add page break after each node (except last)
                if node_name != sorted(grouped_logs.keys())[-1]:
                    story.append(PageBreak())
            
            doc.build(story)

        except Exception as e:
            raise RuntimeError(f"PDF generation failed: {str(e)}")

    def generate_docx(self, logs: List[Dict], output_path: str, 
                     lines_mode: str = 'all',
                     line_limit: int = 0,
                     range_start: int = 0,
                     range_end: int = 0):
        """Generate DOCX with node-based grouping, TOC, and line wrapping"""
        doc = Document()
        
        # Configure styles
        styles = doc.styles
        
        # Create custom styles if they don't exist
        try:
            title_style = styles.add_style('LogTitle', 1)
        except ValueError:
            title_style = styles['LogTitle']
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(14)
        
        body_style = styles['Normal']
        body_style.font.name = 'Courier New'
        body_style.font.size = Pt(10)
        
        # Add title and TOC placeholder
        doc.add_heading('Log Report - Node Overview', 0)
        doc.add_paragraph('Table of Contents', style='Heading 1')
        doc.add_paragraph()  # TOC will be generated when document is opened in Word
        doc.add_page_break()
        
        # Group logs by node
        grouped_logs = self.group_logs_by_node(logs)
        
        # Sort nodes alphabetically
        for node_name in sorted(grouped_logs.keys()):
            # Add node heading (creates TOC entry)
            doc.add_heading(f'Node: {node_name}', level=1)
            
            # Process file types in order (.fbc → .rpc → .log → .lis)
            for file_type in sorted(grouped_logs[node_name].keys(),
                                   key=lambda x: self.TYPE_ORDER.get(x, 999)):
                file_logs = grouped_logs[node_name][file_type]
                
                if not file_logs:  # Skip empty file types
                    continue
                
                for log in file_logs:
                    # Add file subheading
                    doc.add_heading(
                        f'{file_type.upper()[1:]} File: {log["filename"]}', 
                        level=2
                    )
                    
                    # Process lines with filtering
                    content = filter_lines(
                        log['content'],
                        mode=lines_mode,
                        limit=line_limit or log.get('line_limit', 0),
                        start=range_start or log.get('range_start', 0),
                        end=range_end or log.get('range_end', 0)
                    )
                    
                    # Wrap long lines
                    content = self.wrap_long_lines(content)
                    
                    # Combine all lines into a single paragraph with preserved newlines
                    full_text = '\n'.join(content)
                    doc.add_paragraph(full_text, style='Normal')
                    
            doc.add_page_break()
            
        doc.save(output_path)